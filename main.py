# يسطا الكود ده بياخد الصوره ويرفعها على جوجل لينز وياخد النص منها ويضيفها في الوورد
import os
import time
import pyperclip
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ====== Config ======
image_folder = r"images"
output_docx = r"script.docx"
lens_url = "https://lens.google.com/upload"
driver_path = r"msedgedriver.exe"

# ====== Start Edge Browser ======
service = Service(executable_path=driver_path)
options = Options()

# Optional: Add settings to avoid detection
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_argument("--disable-notifications")
options.add_argument("--disable-popup-blocking")
options.add_argument("--disable-infobars")
options.add_argument("--start-maximized")

driver = webdriver.Edge(service=service, options=options)

# ====== Create Word Document ======
doc = Document()
doc.add_heading('Extracted Texts from Google Lens', 0)

# Enable Right-To-Left (RTL) for Arabic text
def set_rtl(paragraph):
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(rtl)

# ====== Loop through images ======
image_files = [img for img in os.listdir(image_folder) if img.lower().endswith(('.jpg', '.png', '.jpeg'))]

for image_name in image_files:
    image_path = os.path.abspath(os.path.join(image_folder, image_name))
    print(f"\nUploading image: {image_name}")

    # Step 1: Open Google Lens upload page
    driver.get(lens_url)
    time.sleep(2)

    # Step 2: Upload image
    try:
        upload_input = WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.XPATH, '//input[@type="file"]'))
        )
        upload_input.send_keys(image_path)
        print("image uploaded")
    except Exception as e:
        print(f"cant upload the photo: {e}")
        continue

    # NEW: Wait 4 seconds to ensure image is fully loaded
    print("waiting 4 seconds for image to load...")
    time.sleep(4)

    # Step 5: Click on "اختيار نص"
    try:
        select_text_button = WebDriverWait(driver, 30).until(
            EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'اختيار نص')]"))
        )
        select_text_button.click()
        print("selecting is done")
    except Exception as e:
        print(f"cant find select text button: {e}")
        continue

    # Step 6: Click on "نسخ"
    try:
        copy_button = WebDriverWait(driver, 30).until(
            EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'نسخ')]"))
        )
        copy_button.click()
        print("copying is done")
    except Exception as e:
        print(f"cant find copy button: {e}")
        continue

    # Read text from clipboard
    extracted_text = pyperclip.paste().strip()

    # Step 7: Write into Word file with image name
    heading = doc.add_heading(f"Image: {image_name}", level=2)
    set_rtl(heading)

    if extracted_text:
        paragraph = doc.add_paragraph(extracted_text)
        set_rtl(paragraph)
    else:
        paragraph = doc.add_paragraph("no text found.")
        set_rtl(paragraph)

    separator = doc.add_paragraph("-" * 40)
    set_rtl(separator)

# Step 8: Save Word file
doc.save(output_docx)
driver.quit()
print(f"\nreslut saved in: {output_docx}")